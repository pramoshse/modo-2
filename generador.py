import streamlit as st
import subprocess
import os

# 1. Instalación inmediata de Playwright
@st.cache_resource
def install_playwright():
    subprocess.run(["playwright", "install", "chromium"])

install_playwright()

import asyncio
from playwright.async_api import async_playwright
import datetime
import base64

import html
import io
import json
import re
from typing import Any, Dict, Iterable, List, Optional

import streamlit.components.v1 as components


# ============================================================
# CONFIGURACIÓN GENERAL
# ============================================================

st.set_page_config(
    page_title="Generador de Procedimientos Modo 2",
    layout="wide",
)

st.markdown(
    """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&display=swap');

        html, body, [class*="css"], .stApp {
            font-family: 'Inter', Arial, sans-serif;
        }

        .main .block-container {
            padding-top: 1.6rem;
            max-width: 1320px;
        }

        .soft-note {
            font-size: 0.78rem;
            color: #64748B;
            line-height: 1.45;
            margin-top: -0.35rem;
            margin-bottom: 1rem;
        }

        .status-ok {
            border: 1px solid #BAE6FD;
            background: #F0F9FF;
            color: #0369A1;
            border-radius: 10px;
            padding: 0.8rem 1rem;
            font-size: 0.9rem;
        }

        .status-warn {
            border: 1px solid #FDE68A;
            background: #FFFBEB;
            color: #92400E;
            border-radius: 10px;
            padding: 0.8rem 1rem;
            font-size: 0.9rem;
        }
    </style>
    """,
    unsafe_allow_html=True,
)

# Color principal Modo 2
MODO1_COLOR = "#00B0F0"  # Modo 2 mantiene el mismo color que Modo 1
MODO1_COLOR_HEX = "00B0F0"       # Sin # para uso en Word/Excel
MODO1_FONT_COLOR = "#FFFFFF"     # Texto blanco sobre fondo azul
MODO1_FONT_HEX = "FFFFFF"        # Sin # para Word/Excel


# ============================================================
# UTILIDADES DE DATOS
# ============================================================

def _html(value: Any) -> str:
    if value is None:
        return ""
    return html.escape(str(value), quote=True)

def _normalize(value: Any) -> str:
    if value is None:
        return ""
    value = str(value).strip()
    value = re.sub(r"\s+", " ", value)
    return value

def _upper(value: Any) -> str:
    return _normalize(value).upper()

def _get_from_dict(source: Dict[str, Any], keys: Iterable[str], default: str = "") -> Any:
    for key in keys:
        if key in source and source.get(key) not in (None, "", [], {}):
            return source.get(key)
    return default

def _get_nested(source: Dict[str, Any], path: Iterable[str], default: Any = "") -> Any:
    current: Any = source
    for key in path:
        if not isinstance(current, dict) or key not in current:
            return default
        current = current[key]
    return current if current not in (None, "") else default

def _split_tasks(raw: Any) -> List[str]:
    if raw is None:
        return []
    if isinstance(raw, dict):
        for key in ("tareas", "items", "seleccionadas", "predefinidas", "values"):
            if key in raw and raw.get(key) not in (None, "", [], {}):
                return _split_tasks(raw.get(key))
        candidate = _get_from_dict(
            raw,
            ["tarea", "actividad", "descripcion", "descripción", "descripcion_tarea", "nombre", "texto", "label"],
            "",
        )
        return [str(candidate).strip(" -•\t")] if str(candidate).strip() else []
    if isinstance(raw, (list, tuple)):
        tasks: List[str] = []
        for item in raw:
            if isinstance(item, dict):
                tasks.extend(_split_tasks(item))
            else:
                item_text = str(item).strip(" -•\t")
                if item_text:
                    tasks.append(item_text)
        return [task for task in tasks if task]
    text = str(raw).strip()
    if not text:
        return []
    parts = []
    for line in re.split(r"[\n\r]+", text):
        line = line.strip(" -•\t")
        if not line:
            continue
        parts.append(line)
    return parts

def _mode_is_modo_2(mode: Any) -> bool:
    clean = _normalize(mode).lower().replace("°", "")
    return clean in {"modo 2", "modo2", "2", "m2"}

def _auto_negocio_from_sitio(sitio: Any, default: str = "") -> str:
    sitio_txt = _normalize(sitio).lower()
    if not sitio_txt:
        return _normalize(default)
    if sitio_txt.startswith("planta") or sitio_txt.startswith("cedi"):
        return "Bebidas"
    return "Ingenios"

def _bytes_to_data_uri(content: bytes, mime: str = "image/png") -> str:
    if not content:
        return ""
    encoded = base64.b64encode(content).decode("utf-8")
    return f"data:{mime or 'image/png'};base64,{encoded}"

def _mime_from_filename(filename: str, default: str = "image/png") -> str:
    ext = os.path.splitext(str(filename or ""))[1].lower()
    return {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
    }.get(ext, default)

def _uploaded_file_data_uri(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    try:
        content = uploaded_file.getvalue()
        mime = getattr(uploaded_file, "type", "") or _mime_from_filename(getattr(uploaded_file, "name", ""))
        return _bytes_to_data_uri(content, mime)
    except Exception:
        return ""

def _local_file_data_uri(candidates: Iterable[str]) -> str:
    for candidate in candidates:
        if not candidate:
            continue
        try:
            if os.path.exists(candidate) and os.path.isfile(candidate):
                with open(candidate, "rb") as fh:
                    return _bytes_to_data_uri(fh.read(), _mime_from_filename(candidate))
        except Exception:
            continue
    return ""

def _logo_data_uri(uploaded_logo=None) -> str:
    uploaded_uri = _uploaded_file_data_uri(uploaded_logo)
    if uploaded_uri:
        return uploaded_uri
    base_dir = os.path.dirname(os.path.abspath(__file__)) if "__file__" in globals() else os.getcwd()
    return _local_file_data_uri(
        [
            os.path.join(base_dir, "arca.png"),
            os.path.join(os.getcwd(), "arca.png"),
            "/home/ubuntu/arca.png",
            "/home/ubuntu/upload/arca.png",
        ]
    )

def _lsr_data_uri() -> str:
    base_dir = os.path.dirname(os.path.abspath(__file__)) if "__file__" in globals() else os.getcwd()
    return _local_file_data_uri(
        [
            os.path.join(base_dir, "LSR.png"),
            os.path.join(base_dir, "lsr.png"),
            os.path.join(os.getcwd(), "LSR.png"),
            os.path.join(os.getcwd(), "lsr.png"),
            "/home/ubuntu/LSR.png",
            "/home/ubuntu/upload/LSR.png",
        ]
    )

def _parse_date_field(value: Any) -> str:
    """
    Normaliza fechas que pueden venir como:
      - str ISO "2026-06-10"  →  "10/06/2026"
      - dict {'__type__': 'date', 'value': '2026-06-10'}  →  "10/06/2026"
      - cualquier otro str  →  se devuelve tal cual (normalizado)
    """
    if value is None:
        return ""
    if isinstance(value, dict):
        raw = value.get("value") or value.get("date") or ""
        value = str(raw).strip()
    value = _normalize(value)
    if not value:
        return ""
    # Intentar parsear ISO YYYY-MM-DD
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})$", value)
    if m:
        return f"{m.group(3)}/{m.group(2)}/{m.group(1)}"
    return value


def extract_procedure_context(payload: Dict[str, Any]) -> Dict[str, Any]:
    widgets = payload.get("widgets") if isinstance(payload.get("widgets"), dict) else {}
    computed = payload.get("computed") if isinstance(payload.get("computed"), dict) else {}

    mode_final = _get_from_dict(
        computed,
        ["modo_final", "modo_convalidado", "modo_resultante"],
        _get_from_dict(widgets, ["modo_final", "modo_convalidado", "modo_resultante"], ""),
    )

    # Combinar tareas_predefinidas + tareas_manuales (ambas en el listado)
    manuales_raw = (_get_from_dict(widgets, ["tareas_manuales"], "")
                    or _get_from_dict(payload, ["tareas_manuales"], ""))
    predefinidas_raw = (
        _get_from_dict(widgets, ["tareas_predefinidas", "tareas_predefinidas_txt", "tareas_preseleccionadas", "tareas_seleccionadas"], "")
        or _get_from_dict(payload, ["tareas_predefinidas", "tareas_predefinidas_txt", "tareas_preseleccionadas", "tareas_seleccionadas"], "")
    )
    fallback_raw = "" if (manuales_raw or predefinidas_raw) else _get_from_dict(
        widgets,
        ["tareas_txt", "tareas", "tareas_aplicables", "descripcion_tareas", "actividad", "actividades"],
        "",
    )
    tasks_manuales = _split_tasks(manuales_raw)
    tasks_predefinidas = _split_tasks(predefinidas_raw)
    tasks_fallback = _split_tasks(fallback_raw)
    # Unir sin duplicados: predefinidas primero, luego manuales
    seen: set = set()
    tasks: List[str] = []
    for t in tasks_predefinidas + tasks_manuales + tasks_fallback:
        if t not in seen:
            seen.add(t)
            tasks.append(t)
    if not tasks:
        tasks = [
            "Ingreso de cuerpo completo al equipo con dos paros de operación independientes aplicados.",
            "Intervención a través de protección no entrelazada con sistema de llave atrapada activo.",
            "Tarea de mantenimiento menor con bloqueo de control de seguridad y candado personal.",
        ]

    energias = _get_from_dict(computed, ["energias_seleccionadas", "energias"], {})
    lista_peligros = _get_from_dict(computed, ["lista_peligros_final", "peligros"], [])
    evaluaciones_fine = _get_from_dict(computed, ["evaluaciones_fine"], [])

    # Código y fecha de la evaluación de riesgos de referencia
    codigo_borrador = _get_from_dict(
        payload,
        ["codigo_documento", "codigo", "codigo_borrador", "id", "referencia"],
        _get_from_dict(widgets, ["codigo_documento", "codigo", "codigo_borrador", "id", "referencia"], ""),
    )
    fecha_raw = _get_from_dict(
        payload,
        ["fecha_documento", "fecha", "fecha_borrador", "fecha_evaluacion"],
        _get_from_dict(widgets, ["fecha_documento", "fecha", "fecha_borrador", "fecha_evaluacion"], ""),
    )
    fecha_borrador = _parse_date_field(fecha_raw)

    sitio = _get_from_dict(widgets, ["sitio", "planta", "site"], "")
    negocio_json = _get_from_dict(widgets, ["negocio", "business", "unidad_negocio"], "")

    return {
        "negocio": _auto_negocio_from_sitio(sitio, negocio_json),
        "sitio": sitio,
        "tipo_sitio": _get_from_dict(widgets, ["tipo_sitio"], ""),
        "area": _get_from_dict(widgets, ["area_sector", "area", "sector", "área"], ""),
        "linea": _get_from_dict(widgets, ["linea", "línea", "linea_equipo"], ""),
        "equipo": _get_from_dict(
            widgets,
            ["equipo_desc", "equipo", "descripcion_equipo", "nombre_equipo", "maquina", "máquina"],
            "EQUIPO SIN DESCRIPCIÓN",
        ),
        "fabricante": _get_from_dict(widgets, ["fabricante", "marca"], ""),
        "modelo": _get_from_dict(widgets, ["modelo"], ""),
        "anio": _get_from_dict(widgets, ["anio", "año"], ""),
        "modo_inicial": _get_from_dict(computed, ["modo_inicial"], _get_from_dict(widgets, ["modo_inicial"], "")),
        "modo_final": mode_final,
        "tareas": tasks,
        "energias": energias,
        "lista_peligros": lista_peligros,
        "evaluaciones_fine": evaluaciones_fine,
        "codigo_documento": codigo_borrador,
        "fecha_documento": fecha_borrador,
        "photo_uri": "",
    }


# ============================================================
# PLANTILLA HTML/PDF MODO 2
# ============================================================

def _tasks_html(tasks: List[str]) -> str:
    if len(tasks) > 13:
        # Modo compacto: todas en una sola línea separadas por " - "
        inline = " - ".join(_html(task) for task in tasks)
        return f"<div class='task-line task-inline'>{inline}</div>"
    return "".join(f"<div class='task-line'>- {_html(task)}</div>" for task in tasks)

def _equipment_meta(ctx: Dict[str, Any]) -> str:
    meta = []
    if ctx.get("fabricante"):
        meta.append(f"Fabricante: {_html(ctx.get('fabricante'))}")
    if ctx.get("modelo"):
        meta.append(f"Modelo: {_html(ctx.get('modelo'))}")
    if ctx.get("anio"):
        meta.append(f"Año: {_html(ctx.get('anio'))}")
    return " · ".join(meta)


# Textos de Acción y Verificación Modo 2
_ACCION_MODO2 = [
    "1. Detener el equipo utilizando el control normal de PARO DE OPERACIÓN.",
    "2. Confirmar que la tarea corresponde a Modo 2: ingreso de cuerpo completo o intervención a través de protecciones no entrelazadas.",
    "3. Aplicar al menos dos funciones independientes de PARO DE OPERACIÓN antes de iniciar la intervención.",
    "4. Bloquear al menos un sistema de control de seguridad para impedir el reinicio inesperado del equipo.",
    "5. Seleccionar el método de control aplicable según la evaluación de riesgos del equipo y de la tarea.",
    "6. Si se utiliza puerta de acceso entrelazada, instalar un dispositivo de bloqueo que impida el cierre de la puerta.",
    "7. Si se utiliza cortina de luz, instalar un dispositivo de bloqueo que garantice que la cortina permanezca activada durante la intervención.",
    "8. Si se utiliza sistema de llave atrapada, retirar la llave única del control de arranque/parada y utilizarla para abrir y bloquear la puerta de acceso en posición segura.",
    "9. Si se utiliza bloqueo de la parada de operación, activar la parada e instalar un dispositivo de bloqueo que impida su restablecimiento.",
    "10. Cada trabajador que ingrese al equipo debe aplicar su propio candado, llave o dispositivo único de control.",
    "11. Cada trabajador debe conservar la posesión exclusiva de su llave durante toda la intervención.",
    "12. Ingresar al equipo únicamente cuando las condiciones de seguridad hayan sido confirmadas.",
    "13. Realizar únicamente la tarea autorizada y definida para Modo 2.",
    "14. Mantener comunicación y control sobre todas las personas involucradas durante la intervención.",
    "15. Retirar herramientas, materiales, residuos y elementos sueltos antes de cerrar o restablecer el equipo.",
    "16. Confirmar que todo el personal haya salido completamente del equipo antes de retirar bloqueos o restablecer condiciones operativas.",
    "17. Retirar los candados o dispositivos personales únicamente por la persona que los instaló.",
    "18. Restablecer el equipo solo cuando todos los dispositivos de seguridad estén nuevamente operativos y el área se encuentre despejada.",
    "19. Reiniciar el equipo únicamente mediante el comando normal de arranque, luego de confirmar condición segura.",
    "20. Si la tarea requiere desmontar guardas, retirar cubiertas o anular protecciones, reclasificar la intervención como Modo 3 / LOTO.",
]

_VERIFICACION_MODO2 = [
    "1. Verificar que la intervención corresponda a una tarea de Modo 2 y no a Modo 1 ni Modo 3.",
    "2. Verificar que la tarea implique ingreso de cuerpo completo al equipo o intervención a través de protecciones no entrelazadas.",
    "3. Verificar que el cuerpo del trabajador no pueda actuar como barrera o bloqueo efectivo frente al reinicio del equipo.",
    "4. Verificar que se hayan aplicado al menos dos funciones independientes de PARO DE OPERACIÓN.",
    "5. Verificar que al menos un sistema de control de seguridad esté bloqueado para impedir el reinicio.",
    "6. Verificar que cada trabajador involucrado cuente con su propia llave, candado o dispositivo único.",
    "7. Verificar que ningún trabajador dependa del candado, llave o dispositivo de otra persona.",
    "8. Verificar que los dispositivos de seguridad aplicables (interlocks o cortinas de luz) cumplan como mínimo con PLr = d / ANSI Cat. 3.",
    "9. Verificar que, si los dispositivos de seguridad no cumplen con PLr = d / ANSI Cat. 3, la tarea sea ejecutada bajo LOTO.",
    "10. Verificar que el paro utilizado sea el PARO DE OPERACIÓN normal y no el paro de emergencia como método habitual de control.",
    "11. Verificar que la puerta de acceso entrelazada, si aplica, quede bloqueada en condición abierta o impedida de cierre.",
    "12. Verificar que cada trabajador que ingresa haya aplicado su candado al dispositivo de bloqueo de la puerta.",
    "13. Verificar que la cortina de luz, si aplica, permanezca activada y bloqueada durante toda la intervención.",
    "14. Verificar que cada trabajador que ingresa haya aplicado su candado al dispositivo de bloqueo asociado a la cortina de luz.",
    "15. Verificar que los sistemas de llave atrapada, si aplican, impidan el arranque mientras la llave no se encuentre restituida al control correspondiente.",
    "16. Verificar que cada trabajador conserve la llave asignada mientras permanezca dentro del equipo.",
    "17. Verificar que el bloqueo de la parada de operación, si aplica, impida su restablecimiento no autorizado.",
    "18. Verificar que no existan bypass, anulaciones, puentes, manipulaciones o condiciones inseguras en los dispositivos de seguridad.",
    "19. Verificar que no se retiren guardas, cubiertas o protecciones; si se requiere retirarlas, reclasificar como Modo 3 / LOTO.",
    "20. Verificar que no exista exposición a energías peligrosas no controladas durante la intervención.",
    "21. Verificar que todos los trabajadores hayan salido completamente del equipo antes de retirar bloqueos.",
    "22. Verificar que todas las herramientas, piezas, materiales o residuos hayan sido retirados del interior del equipo.",
    "23. Verificar que los dispositivos de seguridad queden restablecidos, operativos y sin fallas antes del reinicio.",
    "24. Verificar que el equipo no pueda reiniciar automáticamente al cerrar una puerta, restablecer una cortina o retirar un bloqueo.",
    "25. Verificar que el reinicio requiera una acción voluntaria, independiente y controlada del operador.",
]


def build_modo_2_html(
    ctx: Dict[str, Any],
    *,
    codigo: str,
    revision: str,
    fecha: datetime.date,
    organizacion: str,
    logo_uri: str = "",
    lsr_uri: str = "",
    personal_afectado: str,
    personal_autorizado: str,
    elaborado_por: str,
    aprobado_por: str,
    puesto_elaborado: str = "",
    puesto_aprobado: str = "",
    fecha_firma: "datetime.date | str | None" = None,
    eval_riesgos_codigo: str = "",
    eval_riesgos_fecha: str = "",
) -> str:
    fecha_txt = fecha.strftime("%d/%m/%Y") if isinstance(fecha, datetime.date) else _normalize(fecha)
    fecha_firma_txt = fecha_firma.strftime("%d/%m/%Y") if isinstance(fecha_firma, datetime.date) else (_normalize(fecha_firma) or fecha_txt)
    task_rows = _tasks_html(ctx.get("tareas") or [])
    photo_uri = ctx.get("photo_uri", "")
    equipo_meta = _equipment_meta(ctx)

    logo_block = (
        f"<img class='logo-img' src='{logo_uri}' alt='Logo'>"
        if logo_uri
        else f"<div class='logo-text'>{_html(organizacion)}</div>"
    )

    photo_block = (
        f"<img class='equipment-photo' src='{photo_uri}' alt='Foto del paso a paso'>"
        if photo_uri
        else "<div class='photo-placeholder'>FOTO DEL PASO A PASO</div>"
    )

    lsr_block = (
        f"<img class='lsr-img' src='{lsr_uri}' alt='LSR'>"
        if lsr_uri
        else "<div class='lsr-placeholder'>LSR</div>"
    )

    # Evaluación de riesgos text
    eval_ref = ""
    if eval_riesgos_codigo:
        eval_ref += f"Código: {_html(eval_riesgos_codigo)}"
    if eval_riesgos_fecha:
        eval_ref += f"  ·  Fecha: {_html(eval_riesgos_fecha)}"
    if not eval_ref:
        eval_ref = "—"

    accion_html = "".join(f"<p>{_html(line)}</p>" for line in _ACCION_MODO2)
    verif_html = "".join(f"<p>{_html(line)}</p>" for line in _VERIFICACION_MODO2)

    html_doc = f"""
<!doctype html>
<html lang="es">
<head>
<meta charset="utf-8">
<title>Procedimiento Modo 2 - Control de Energías Peligrosas</title>
<style>
    @page {{
        size: A4 portrait;
        margin: 5mm;
    }}

    * {{
        box-sizing: border-box;
    }}

    body {{
        margin: 0;
        background: #E5E7EB;
        color: #0F172A;
        font-family: Bahnschrift, 'Bahnschrift SemiCondensed', 'Arial Narrow', Arial, Helvetica, sans-serif;
        font-size: 8.7px;
        line-height: 1.14;
    }}

    .sheet {{
        width: 760px;
        min-height: 0;
        margin: 0 auto;
        background: #FFFFFF;
        border: 1.2px solid #111827;
        box-shadow: 0 12px 28px rgba(15, 23, 42, 0.18);
    }}

    table {{
        border-collapse: collapse;
        width: 100%;
        table-layout: fixed;
    }}

    td, th {{
        border: 1px solid #111827;
        padding: 5px 6px;
        vertical-align: middle;
    }}

    .top-black {{
        background: #050505;
        color: #FFFFFF;
        font-weight: 800;
        text-align: center;
        letter-spacing: 0.06em;
        font-size: 11px;
        padding: 2px 4px;
        text-transform: uppercase;
    }}

    .logo-cell {{
        width: 62px;
        text-align: center;
        background: #FFFFFF;
        padding: 3px;
    }}

    .logo-img {{
        max-width: 56px;
        max-height: 40px;
        object-fit: contain;
    }}

    .logo-text {{
        color: #B91C1C;
        font-size: 9px;
        font-weight: 900;
        line-height: 1.05;
        text-transform: uppercase;
        word-break: break-word;
    }}

    .main-title {{
        text-align: center;
        font-size: 12.5px;
        font-weight: 900;
        letter-spacing: 0.01em;
        text-transform: uppercase;
        padding: 5px 6px;
    }}

    .doc-data-label {{
        background: #E5E7EB;
        width: 68px;
        text-align: center;
        font-weight: 800;
        text-transform: uppercase;
        font-size: 10px;
    }}

    .doc-data-value {{
        width: 86px;
        text-align: center;
        font-weight: 700;
        background: #F8FAFC;
        font-size: 10px;
    }}

    .info-label {{
        display: block;
        color: #334155;
        font-size: 10px;
        margin-bottom: 2px;
    }}

    .info-value {{
        display: block;
        color: #111827;
        font-size: 13px;
        font-weight: 700;
    }}

    .info-box {{
        height: 50px;
        text-align: center;
        background: #FFFFFF;
    }}

    .person-title {{
        font-weight: 900;
        text-align: center;
        background: {MODO1_COLOR};
        color: #FFFFFF;
        font-size: 8.8px;
        padding: 3px 6px;
    }}

    .person-body {{
        text-align: center;
        min-height: 28px;
        font-size: 8.8px;
        white-space: pre-line;
    }}

    .equipment-label {{
        width: 62px;
        font-weight: 800;
        color: #334155;
        background: #F8FAFC;
        text-align: center;
    }}

    .equipment-title {{
        text-align: center;
        font-size: 15px;
        font-weight: 900;
        text-transform: uppercase;
        padding: 4px 6px 2px 6px;
    }}

    .equipment-meta {{
        display: block;
        margin-top: 3px;
        color: #64748B;
        font-size: 9px;
        font-weight: 600;
        text-transform: none;
    }}

    /* Barra de evaluación de riesgos */
    .eval-bar {{
        background: #FFFFFF;
        color: #0F172A;
        font-weight: 700;
        font-size: 8.5px;
        padding: 4px 8px;
        border-left: 1px solid #111827;
        border-right: 1px solid #111827;
        border-bottom: 1px solid #111827;
    }}

    .red-header th,
    .red-bar {{
        background: {MODO1_COLOR};
        color: #FFFFFF;
        font-weight: 900;
        text-align: center;
        font-size: 10px;
    }}

    .red-bar {{
        border-left: 1px solid #111827;
        border-right: 1px solid #111827;
        padding: 4px 8px;
        text-transform: none;
    }}

    .block-two {{
        text-align: center;
        font-size: 48px;
        font-weight: 900;
        color: #000000;
        background: #FFFFFF;
        line-height: 1;
    }}

    .mode-box {{
        background: {MODO1_COLOR};
        color: #FFFFFF;
        text-align: center;
        font-size: 17px;
        font-weight: 900;
        text-transform: uppercase;
    }}

    .tasks-cell {{
        min-height: 36px;
        font-size: 8.6px;
        background: #FFFFFF;
        padding: 5px 8px;
    }}

    .task-line {{
        margin: 0 0 4px 0;
    }}

    .task-inline {{
        margin: 0;
        line-height: 1.35;
        white-space: normal;
        word-break: break-word;
    }}

    .photo-area {{
        height: 210px;
        background: #FFFFFF;
        text-align: center;
        vertical-align: middle;
        padding: 7px;
    }}

    .equipment-photo {{
        max-width: 100%;
        max-height: 195px;
        object-fit: contain;
        border: 1px solid #CBD5E1;
        background: #F8FAFC;
        padding: 3px;
    }}

    .photo-placeholder {{
        height: 190px;
        display: flex;
        align-items: center;
        justify-content: center;
        color: #94A3B8;
        font-size: 14px;
        font-weight: 800;
        letter-spacing: 0.03em;
        border: 1px dashed #CBD5E1;
        background: #F8FAFC;
    }}

    .dark-head th {{
        background: #3B3B3B;
        color: #FFFFFF;
        font-weight: 900;
        text-align: center;
        font-size: 8.3px;
        padding: 3px 2px;
    }}

    .procedure-table td {{
        font-size: 7.8px;
        text-align: left;
        background: #FFFFFF;
        vertical-align: top;
        padding: 7px 9px;
    }}

    .procedure-note {{
        font-size: 8.4px;
        line-height: 1.42;
    }}

    .procedure-note p {{
        margin: 0 0 4px 0;
    }}

    .procedure-note p:last-child {{
        margin-bottom: 0;
    }}

    .legend-title {{
        background: {MODO1_COLOR};
        color: #FFFFFF;
        font-weight: 900;
        text-align: center;
        padding: 2px;
    }}

    .energy-legend td,
    .lock-legend td {{
        text-align: center;
        font-weight: 900;
        font-size: 7.6px;
        color: #0F172A;
        padding: 3px 2px;
    }}

    .e-electric {{ background: #000000; color: #FFFFFF !important; }}
    .e-neumatic {{ background: #0284C7; color: #FFFFFF !important; }}
    .e-amoniaco {{
        background: linear-gradient(90deg,
            #D9D9D9 0% 15%, #F59E0B 15% 26%,
            #D9D9D9 26% 42%, #F59E0B 42% 53%,
            #D9D9D9 53% 68%, #F59E0B 68% 79%,
            #D9D9D9 79% 100%);
    }}
    .e-termica {{ background: #DC2626; color: #FFFFFF !important; }}
    .e-hidraulica {{ background: #7C3AED; color: #FFFFFF !important; }}
    .e-potencial {{
        background: linear-gradient(90deg,
            #FFF200 0% 22%, #050505 22% 34%,
            #FFF200 34% 62%, #050505 62% 74%,
            #FFF200 74% 100%);
    }}
    .e-quimica {{ background: #FFF200; }}
    .e-vapor {{ background: #F59E0B; color: #111827 !important; }}
    .e-agua {{ background: #16A34A; color: #FFFFFF !important; }}
    .e-soda {{
        background: linear-gradient(90deg,
            #D9D9D9 0% 25%, #F59E0B 25% 37%,
            #D9D9D9 37% 60%, #F59E0B 60% 72%,
            #D9D9D9 72% 100%);
    }}
    .e-ozono {{ background: #BAE6FD; }}
    .e-gas {{ background: #C7D2FE; color: #111827 !important; }}

    .l-mmto {{ background: #EF0000; color: #FFFFFF !important; }}
    .l-calidad {{ background: #FFF200; }}
    .l-produccion {{ background: #16A34A; color: #FFFFFF !important; }}
    .l-edilicio {{ background: #0F7DBD; color: #FFFFFF !important; }}
    .l-supervisor {{ background: #050505; color: #FFFFFF !important; }}

    .lock-legend {{ table-layout: fixed; }}
    .lock-legend td {{
        width: 20%;
        height: 24px;
        line-height: 1.2;
        vertical-align: middle;
        text-align: center;
        font-weight: 900;
        font-size: 7.6px;
        overflow: hidden;
    }}

    .lsr-section {{
        border-left: 1px solid #111827;
        border-right: 1px solid #111827;
        border-bottom: 1px solid #111827;
        text-align: center;
        padding: 8px 6px;
        background: #FFFFFF;
    }}

    .lsr-img {{
        max-width: 100%;
        max-height: 160px;
        object-fit: contain;
    }}

    .lsr-placeholder {{
        height: 80px;
        display: flex;
        align-items: center;
        justify-content: center;
        color: #94A3B8;
        font-size: 13px;
        font-weight: 800;
        letter-spacing: 0.05em;
        border: 1px dashed #CBD5E1;
        background: #F8FAFC;
    }}

    .footer-sign {{
        margin-top: 0;
    }}

    .signature-cell {{
        height: 30px;
        font-size: 7.8px;
        color: #334155;
        background: #F8FAFC;
        text-align: center;
    }}

    .small-muted {{
        color: #64748B;
        font-size: 8.8px;
        font-weight: 600;
    }}

    @media print {{
        body {{ background: #FFFFFF; }}
        .sheet {{ box-shadow: none; margin: 0; width: 100%; }}
    }}
</style>
</head>
<body>
<div class="sheet">

    <!-- ENCABEZADO -->
    <table>
        <tr>
            <td class="logo-cell" rowspan="3">{logo_block}</td>
            <td class="top-black" colspan="2">CONTROL DE ENERGÍAS PELIGROSAS</td>
            <td class="doc-data-label">Código</td>
            <td class="doc-data-value">{_html(codigo)}</td>
        </tr>
        <tr>
            <td class="main-title" colspan="2" rowspan="2">PROCEDIMIENTO ESPECÍFICO PARA CONTROL DE ENERGÍAS</td>
            <td class="doc-data-label">Revisión</td>
            <td class="doc-data-value">{_html(revision)}</td>
        </tr>
        <tr>
            <td class="doc-data-label">Fecha</td>
            <td class="doc-data-value">{_html(fecha_txt)}</td>
        </tr>
    </table>

    <!-- INFO SITIO + PERSONAL -->
    <table>
        <tr>
            <td class="info-box" style="width: 62px;">
                <span class="info-label">Negocio:</span>
                <span class="info-value">{_html(ctx.get('negocio') or '-')}</span>
            </td>
            <td class="info-box" style="width: 86px;">
                <span class="info-label">Sitio:</span>
                <span class="info-value">{_html(ctx.get('sitio') or '-')}</span>
            </td>
            <td class="info-box" style="width: 86px;">
                <span class="info-label">Área:</span>
                <span class="info-value">{_html(ctx.get('area') or '-')}</span>
            </td>
            <td class="info-box" style="width: 112px;">
                <span class="info-label">Línea:</span>
                <span class="info-value">{_html(ctx.get('linea') or '-')}</span>
            </td>
            <td style="padding:0;">
                <table>
                    <tr><td class="person-title">Personal afectado - Puestos de trabajo</td></tr>
                    <tr><td class="person-body">{_html(personal_afectado)}</td></tr>
                    <tr><td class="person-title">Personal autorizado - Puestos de trabajo</td></tr>
                    <tr><td class="person-body">{_html(personal_autorizado)}</td></tr>
                </table>
            </td>
        </tr>
    </table>

    <!-- EQUIPO -->
    <table>
        <tr>
            <td class="equipment-label">Equipo:</td>
            <td class="equipment-title">
                {_html(ctx.get('equipo') or 'EQUIPO')}
                <span class="equipment-meta">{equipo_meta}</span>
            </td>
        </tr>
    </table>

    <!-- EVALUACIÓN DE RIESGOS -->
    <div class="eval-bar">
        <strong>Evaluación de riesgos de referencia:</strong>&nbsp;&nbsp;{eval_ref}
    </div>

    <!-- TAREAS -->
    <table>
        <tr class="red-header">
            <th style="width: 62px;">Puntos de<br>Bloqueo</th>
            <th style="width: 86px;">Modo de<br>Intervención</th>
            <th>Listado de tareas aplicable al presente procedimiento</th>
        </tr>
        <tr>
            <td class="block-two">0</td>
            <td class="mode-box">MODO 2</td>
            <td class="tasks-cell">{task_rows}</td>
        </tr>
    </table>

    <!-- BARRA SECCIÓN -->
    <div class="red-bar">Procedimiento - Control de Energías Peligrosas</div>

    <!-- FOTO -->
    <table>
        <tr>
            <td class="photo-area">{photo_block}</td>
        </tr>
    </table>

    <!-- PROCEDIMIENTO: Acción y Verificación -->
    <table class="procedure-table">
        <tr class="dark-head">
            <th style="width: 50%;">Acción</th>
            <th style="width: 50%;">Verificación</th>
        </tr>
        <tr>
            <td class="procedure-note">{accion_html}</td>
            <td class="procedure-note">{verif_html}</td>
        </tr>
    </table>

    <!-- LEYENDA ENERGÍAS -->
    <div class="legend-title">Clasificación de Energías Peligrosas</div>
    <table class="energy-legend">
        <tr>
            <td class="e-electric">E: Eléctrica</td>
            <td class="e-neumatic">N: Neumática</td>
            <td class="e-amoniaco">AM: Amoníaco</td>
            <td class="e-termica">T: Térmica</td>
            <td class="e-hidraulica">H: Hidráulica</td>
            <td class="e-potencial">P: Potencial</td>
        </tr>
        <tr>
            <td class="e-quimica">Q: Química</td>
            <td class="e-vapor">V: Vapor</td>
            <td class="e-agua">A: Agua</td>
            <td class="e-soda">SC: Soda Cáustica</td>
            <td class="e-ozono">Oz: Ozono</td>
            <td class="e-gas">GC: Gas Carbónico</td>
        </tr>
    </table>

    <!-- LEYENDA CANDADOS -->
    <div class="legend-title">Clasificación de Candados según sector y función</div>
    <table class="lock-legend">
        <tr>
            <td class="l-mmto">Mantenimiento<br>Industrial</td>
            <td class="l-calidad">Calidad</td>
            <td class="l-produccion">Producción</td>
            <td class="l-edilicio">Mantenimiento Edilicio<br>Contratistas</td>
            <td class="l-supervisor">Supervisor MMTO Industrial<br>(Bloqueo Departamental)</td>
        </tr>
    </table>

    <!-- IMAGEN LSR -->
    <div class="lsr-section">
        {lsr_block}
    </div>

    <!-- FIRMAS -->
    <table class="footer-sign">
        <tr>
            <td class="signature-cell">
                Elaborado por: <strong>{_html(elaborado_por or '-')}</strong><br>
                <span class="small-muted">Puesto: {_html(puesto_elaborado or '-')} · Fecha: {_html(fecha_firma_txt)}</span>
            </td>
            <td class="signature-cell">
                Aprobado por: <strong>{_html(aprobado_por or '-')}</strong><br>
                <span class="small-muted">Puesto: {_html(puesto_aprobado or '-')} · Fecha: {_html(fecha_firma_txt)}</span>
            </td>
        </tr>
    </table>

</div>
</body>
</html>
"""
    return html_doc


# ============================================================
# EXPORTACIÓN PDF
# ============================================================

def _get_playwright_sync_api():
    try:
        module = __import__("playwright.sync_api", fromlist=["sync_playwright"])
        return getattr(module, "sync_playwright", None)
    except Exception:
        return None

def html_to_pdf_bytes(html_doc: str) -> bytes:
    sync_playwright = _get_playwright_sync_api()
    if sync_playwright is None:
        raise RuntimeError(
            "La exportación PDF requiere Playwright. Instalá la dependencia con: pip install playwright && python -m playwright install chromium."
        )
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page(viewport={"width": 1240, "height": 1754})
            page.set_content(html_doc, wait_until="networkidle")
            page.emulate_media(media="print")
            pdf_bytes = page.pdf(
                format="A4",
                landscape=False,
                scale=0.90,
                print_background=True,
                prefer_css_page_size=True,
                margin={"top": "0mm", "right": "0mm", "bottom": "0mm", "left": "0mm"},
            )
            browser.close()
            return pdf_bytes
    except Exception as exc:
        raise RuntimeError(f"No se pudo generar el PDF con Playwright. Detalle: {exc}") from exc


# ============================================================
# EXPORTACIÓN WORD
# ============================================================

def _data_uri_to_bytes(data_uri: str) -> "tuple[bytes, str]":
    if not data_uri or not isinstance(data_uri, str) or not data_uri.startswith("data:"):
        return b"", ".png"
    header, _, encoded = data_uri.partition(",")
    if not encoded:
        return b"", ".png"
    mime = header.split(";", 1)[0].replace("data:", "").strip().lower()
    ext = {"image/png": ".png", "image/jpeg": ".jpg", "image/jpg": ".jpg", "image/webp": ".webp"}.get(mime, ".png")
    try:
        return base64.b64decode(encoded), ext
    except Exception:
        return b"", ext

def _fecha_export_txt(value: Any) -> str:
    if isinstance(value, datetime.date):
        return value.strftime("%d/%m/%Y")
    return _normalize(value)

def _docx_set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", "").upper())

def _docx_set_cell_borders(cell, color: str = "111827", size: str = "6") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color.replace("#", "").upper())

def _docx_set_cell_margins(cell, top: int = 45, start: int = 45, bottom: int = 45, end: int = 45) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def _docx_set_cell_width(cell, width_cm: float) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")

def _docx_clear_cell(cell) -> None:
    cell.text = ""

def _docx_write_cell(
    cell,
    text: Any = "",
    *,
    bold: bool = False,
    size: float = 7.5,
    color: str = "0F172A",
    fill: "str | None" = None,
    align: str = "center",
    valign: str = "center",
) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _docx_clear_cell(cell)
    if fill:
        _docx_set_cell_shading(cell, fill)
    _docx_set_cell_borders(cell)
    _docx_set_cell_margins(cell)
    cell.vertical_alignment = {
        "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
        "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
    }.get(valign, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(str(text or ""))
    run.bold = bold
    run.font.name = "Bahnschrift"
    run.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(size)
    run.font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor.from_string(color.replace("#", "").upper())

def _docx_apply_table_grid(table, widths_cm: "list[float] | None" = None) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.shared import Cm
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            _docx_set_cell_borders(cell)
            _docx_set_cell_margins(cell)
            if widths_cm and idx < len(widths_cm):
                _docx_set_cell_width(cell, widths_cm[idx])
    if widths_cm:
        for idx, width in enumerate(widths_cm):
            try:
                table.columns[idx].width = Cm(width)
            except Exception:
                pass

def html_to_word_bytes(
    ctx: Dict[str, Any],
    *,
    codigo: str,
    revision: str,
    fecha: "datetime.date | str",
    organizacion: str,
    logo_uri: str = "",
    lsr_uri: str = "",
    personal_afectado: str,
    personal_autorizado: str,
    elaborado_por: str,
    aprobado_por: str,
    puesto_elaborado: str = "",
    puesto_aprobado: str = "",
    fecha_firma: "datetime.date | str | None" = None,
    eval_riesgos_codigo: str = "",
    eval_riesgos_fecha: str = "",
) -> bytes:
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    fecha_txt = _fecha_export_txt(fecha)
    fecha_firma_txt = _fecha_export_txt(fecha_firma) or fecha_txt
    equipo_meta = _equipment_meta(ctx)

    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.45)
    section.bottom_margin = Cm(0.45)
    section.left_margin = Cm(0.45)
    section.right_margin = Cm(0.45)

    style = document.styles["Normal"]
    style.font.name = "Bahnschrift"
    style.font.size = Pt(7.5)

    # Encabezado
    header = document.add_table(rows=3, cols=5)
    _docx_apply_table_grid(header, [2.0, 5.0, 5.2, 2.4, 3.2])
    header.cell(0, 0).merge(header.cell(2, 0))
    header.cell(0, 1).merge(header.cell(0, 2))
    header.cell(1, 1).merge(header.cell(2, 2))
    logo_cell = header.cell(0, 0)
    _docx_write_cell(logo_cell, "", fill="FFFFFF")
    logo_bytes, _ = _data_uri_to_bytes(logo_uri)
    if logo_bytes:
        try:
            p = logo_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(logo_bytes), width=Cm(1.55))
        except Exception:
            _docx_write_cell(logo_cell, organizacion, bold=True, size=7.0, color="B91C1C", fill="FFFFFF")
    else:
        _docx_write_cell(logo_cell, organizacion, bold=True, size=7.0, color="B91C1C", fill="FFFFFF")
    _docx_write_cell(header.cell(0, 1), "CONTROL DE ENERGÍAS PELIGROSAS", bold=True, size=8.5, color="FFFFFF", fill="050505")
    _docx_write_cell(header.cell(1, 1), "PROCEDIMIENTO ESPECÍFICO PARA CONTROL DE ENERGÍAS", bold=True, size=9.0, fill="FFFFFF")
    for row_idx, label, value in ((0, "Código", codigo), (1, "Revisión", revision), (2, "Fecha", fecha_txt)):
        _docx_write_cell(header.cell(row_idx, 3), label, bold=True, size=7.5, fill="E5E7EB")
        _docx_write_cell(header.cell(row_idx, 4), value, bold=True, size=7.5, fill="F8FAFC")

    # Info sitio + personal
    info = document.add_table(rows=1, cols=5)
    _docx_apply_table_grid(info, [2.0, 2.8, 2.8, 3.6, 6.6])
    for idx, (label, value) in enumerate(
        [("Negocio:", ctx.get("negocio") or "-"), ("Sitio:", ctx.get("sitio") or "-"),
         ("Área:", ctx.get("area") or "-"), ("Línea:", ctx.get("linea") or "-")]
    ):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        cell = info.cell(0, idx)
        _docx_clear_cell(cell)
        _docx_set_cell_shading(cell, "FFFFFF")
        _docx_set_cell_borders(cell)
        _docx_set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = 0
        para.paragraph_format.space_after = 0
        run_label = para.add_run(label + "\n")
        run_label.bold = False
        run_label.font.name = "Bahnschrift"
        run_label.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(7.2)
        run_label.font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor.from_string("334155")
        run_val = para.add_run(str(value))
        run_val.bold = True
        run_val.font.name = "Bahnschrift"
        run_val.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(9.0)
        run_val.font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor.from_string("111827")
    person_cell = info.cell(0, 4)
    _docx_clear_cell(person_cell)
    _docx_set_cell_borders(person_cell)
    _docx_set_cell_margins(person_cell, top=0, start=0, bottom=0, end=0)
    from docx.oxml.ns import qn as _qn
    for p_elem in person_cell._tc.findall(_qn("w:p")):
        person_cell._tc.remove(p_elem)
    nested = person_cell.add_table(rows=4, cols=1)
    _docx_apply_table_grid(nested, [6.4])
    _docx_write_cell(nested.cell(0, 0), "Personal afectado - Puestos de trabajo", bold=True, size=6.8, color=MODO1_FONT_HEX, fill=MODO1_COLOR_HEX)
    _docx_write_cell(nested.cell(1, 0), personal_afectado, size=6.8, fill="FFFFFF")
    _docx_write_cell(nested.cell(2, 0), "Personal autorizado - Puestos de trabajo", bold=True, size=6.8, color=MODO1_FONT_HEX, fill=MODO1_COLOR_HEX)
    _docx_write_cell(nested.cell(3, 0), personal_autorizado, size=6.8, fill="FFFFFF")

    # Equipo
    equipment = document.add_table(rows=1, cols=2)
    _docx_apply_table_grid(equipment, [2.0, 15.8])
    _docx_write_cell(equipment.cell(0, 0), "Equipo:", bold=True, size=7.4, fill="F8FAFC")
    _docx_write_cell(equipment.cell(0, 1), f"{_upper(ctx.get('equipo') or 'EQUIPO')}\n{equipo_meta}", bold=True, size=9.0, fill="FFFFFF")

    # Evaluación de riesgos
    eval_bar = document.add_table(rows=1, cols=1)
    _docx_apply_table_grid(eval_bar, [17.8])
    eval_ref_txt = "Evaluación de riesgos de referencia:  "
    if eval_riesgos_codigo:
        eval_ref_txt += f"Código: {eval_riesgos_codigo}"
    if eval_riesgos_fecha:
        eval_ref_txt += f"  ·  Fecha: {eval_riesgos_fecha}"
    if eval_ref_txt.strip().endswith(":"):
        eval_ref_txt += " —"
    _docx_write_cell(eval_bar.cell(0, 0), eval_ref_txt, bold=True, size=7.0, color="0F172A", fill="FFFFFF", align="left")

    # Tareas
    tasks_table = document.add_table(rows=2, cols=3)
    _docx_apply_table_grid(tasks_table, [2.0, 2.8, 13.0])
    for idx, title in enumerate(["Puntos de\nBloqueo", "Modo de\nIntervención", "Listado de tareas aplicable al presente procedimiento"]):
        _docx_write_cell(tasks_table.cell(0, idx), title, bold=True, size=7.2, color=MODO1_FONT_HEX, fill=MODO1_COLOR_HEX)
    _docx_write_cell(tasks_table.cell(1, 0), "2", bold=True, size=32.0, fill="FFFFFF")
    _docx_write_cell(tasks_table.cell(1, 1), "MODO 2", bold=True, size=13.0, color=MODO1_FONT_HEX, fill=MODO1_COLOR_HEX)
    tasks_text = "\n".join(f"- {task}" for task in (ctx.get("tareas") or []))
    _docx_write_cell(tasks_table.cell(1, 2), tasks_text, size=6.8, fill="FFFFFF", align="left")

    # Barra
    bar = document.add_table(rows=1, cols=1)
    _docx_apply_table_grid(bar, [17.8])
    _docx_write_cell(bar.cell(0, 0), "Procedimiento - Control de Energías Peligrosas", bold=True, size=7.6, color=MODO1_FONT_HEX, fill=MODO1_COLOR_HEX)

    # Foto
    photo_table = document.add_table(rows=1, cols=1)
    _docx_apply_table_grid(photo_table, [17.8])
    photo_cell = photo_table.cell(0, 0)
    _docx_write_cell(photo_cell, "FOTO DEL PASO A PASO", bold=True, size=11.0, color="94A3B8", fill="FFFFFF")
    photo_bytes, _ = _data_uri_to_bytes(ctx.get("photo_uri", ""))
    if photo_bytes:
        try:
            _docx_clear_cell(photo_cell)
            p = photo_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(photo_bytes), width=Cm(12.0))
            _docx_set_cell_borders(photo_cell)
        except Exception:
            _docx_write_cell(photo_cell, "FOTO DEL PASO A PASO", bold=True, size=11.0, color="94A3B8", fill="FFFFFF")

    # Procedimiento: Acción y Verificación
    proc = document.add_table(rows=2, cols=2)
    _docx_apply_table_grid(proc, [8.9, 8.9])
    _docx_write_cell(proc.cell(0, 0), "Acción", bold=True, size=7.0, color="FFFFFF", fill="3B3B3B")
    _docx_write_cell(proc.cell(0, 1), "Verificación", bold=True, size=7.0, color="FFFFFF", fill="3B3B3B")
    accion_txt = "\n\n".join(_ACCION_MODO2)
    verif_txt = "\n\n".join(_VERIFICACION_MODO2)
    _docx_write_cell(proc.cell(1, 0), accion_txt, size=6.5, fill="FFFFFF", align="left", valign="top")
    _docx_write_cell(proc.cell(1, 1), verif_txt, size=6.5, fill="FFFFFF", align="left", valign="top")

    # Leyenda energías
    legend_title = document.add_table(rows=1, cols=1)
    _docx_apply_table_grid(legend_title, [17.8])
    _docx_write_cell(legend_title.cell(0, 0), "Clasificación de Energías Peligrosas", bold=True, size=7.0, color=MODO1_FONT_HEX, fill=MODO1_COLOR_HEX)
    energy = document.add_table(rows=2, cols=6)
    _docx_apply_table_grid(energy, [17.8 / 6] * 6)

    def _make_stripe_png_exact(bg, stripe_color, stripe_positions, width=120, height=30):
        try:
            from PIL import Image as _PIL, ImageDraw as _Draw
            img = _PIL.new("RGB", (width, height), tuple(int(bg[i:i+2], 16) for i in (0, 2, 4)))
            draw = _Draw.Draw(img)
            sc = tuple(int(stripe_color[i:i+2], 16) for i in (0, 2, 4))
            for x0_pct, x1_pct in stripe_positions:
                x0 = int(width * x0_pct / 100)
                x1 = int(width * x1_pct / 100)
                draw.rectangle([x0, 0, x1 - 1, height - 1], fill=sc)
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            return buf.getvalue()
        except Exception:
            return b""

    def _docx_energy_stripe_cell(cell, label, font_color, png_bytes, cell_width_cm, cell_height_cm=0.7):
        from lxml import etree as _etree
        from docx.oxml import OxmlElement as _OE
        from docx.oxml.ns import qn as _qn
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if not png_bytes:
            _docx_write_cell(cell, label, bold=True, size=5.8, color=font_color, fill="D9D9D9")
            return
        cell.text = ""
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_b = tc_pr.first_child_found_in("w:tcBorders")
        if tc_b is None:
            tc_b = _OE("w:tcBorders"); tc_pr.append(tc_b)
        for edge in ("top", "left", "bottom", "right"):
            s = _OE(f"w:{edge}")
            s.set(_qn("w:val"), "single"); s.set(_qn("w:sz"), "6")
            s.set(_qn("w:space"), "0"); s.set(_qn("w:color"), "111827")
            tc_b.append(s)
        tc_m = tc_pr.first_child_found_in("w:tcMar")
        if tc_m is None:
            tc_m = _OE("w:tcMar"); tc_pr.append(tc_m)
        for m in ("top", "start", "bottom", "end"):
            n = _OE(f"w:{m}"); n.set(_qn("w:w"), "0"); n.set(_qn("w:type"), "dxa")
            tc_m.append(n)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        rId, _ = cell.part.get_or_add_image(io.BytesIO(png_bytes))
        cx = int(cell_width_cm * 914400 / 2.54)
        cy = int(cell_height_cm * 914400 / 2.54)
        anchor_xml = (
            f'<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
            f' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
            f' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'
            f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f'<wp:anchor behindDoc="1" distT="0" distB="0" distL="0" distR="0"'
            f' simplePos="0" locked="1" layoutInCell="1" allowOverlap="0" relativeHeight="1">'
            f'<wp:simplePos x="0" y="0"/>'
            f'<wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH>'
            f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>-82400</wp:posOffset></wp:positionV>'
            f'<wp:extent cx="{cx}" cy="{cy}"/>'
            f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
            f'<wp:wrapNone/>'
            f'<wp:docPr id="10" name="bg"/>'
            f'<wp:cNvGraphicFramePr/>'
            f'<a:graphic>'
            f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
            f'<pic:pic>'
            f'<pic:nvPicPr><pic:cNvPr id="10" name="bg"/><pic:cNvPicPr/></pic:nvPicPr>'
            f'<pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
            f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
            f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
            f'</pic:pic></a:graphicData></a:graphic></wp:anchor></w:drawing>'
        )
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = 0
        para.paragraph_format.space_after = 0
        drawing_el = _etree.fromstring(anchor_xml.encode())
        r0 = _OE("w:r"); r0.append(drawing_el); para._p.append(r0)
        r1 = _OE("w:r")
        rPr = _OE("w:rPr")
        fc = font_color.replace("#", "").upper()
        for tag, attrs in [
            ("w:b", {}),
            ("w:rFonts", {_qn("w:ascii"): "Bahnschrift", _qn("w:hAnsi"): "Bahnschrift"}),
            ("w:sz", {_qn("w:val"): "12"}),
            ("w:szCs", {_qn("w:val"): "12"}),
            ("w:color", {_qn("w:val"): fc}),
        ]:
            e = _OE(tag)
            for k, v in attrs.items(): e.set(k, v)
            rPr.append(e)
        r1.append(rPr)
        t_el = _OE("w:t"); t_el.text = label; r1.append(t_el)
        para._p.append(r1)

    _CELL_W_CM = 17.8 / 6
    _CELL_H_CM = 0.65
    energy_items = [
        (0, 0, "E: Eléctrica",    "000000", "FFFFFF", None),
        (0, 1, "N: Neumática",    "0284C7", "FFFFFF", None),
        (0, 2, "AM: Amoníaco",    "D9D9D9", "0F172A", ("D9D9D9", "F59E0B", [(15, 26), (42, 53), (68, 79)])),
        (0, 3, "T: Térmica",      "DC2626", "FFFFFF", None),
        (0, 4, "H: Hidráulica",   "7C3AED", "FFFFFF", None),
        (0, 5, "P: Potencial",    "FFF200", "555555", ("FFF200", "050505", [(22, 34), (62, 74)])),
        (1, 0, "Q: Química",      "FFF200", "0F172A", None),
        (1, 1, "V: Vapor",        "F59E0B", "111827", None),
        (1, 2, "A: Agua",         "16A34A", "FFFFFF", None),
        (1, 3, "SC: Soda Cáustica","D9D9D9","0F172A", ("D9D9D9", "F59E0B", [(25, 37), (60, 72)])),
        (1, 4, "Oz: Ozono",       "BAE6FD", "0F172A", None),
        (1, 5, "GC: Gas Carbónico","C7D2FE","111827", None),
    ]
    for row, col, label, fill, font_color, stripe_spec in energy_items:
        cell = energy.cell(row, col)
        if stripe_spec:
            bg, sc, positions = stripe_spec
            png = _make_stripe_png_exact(bg, sc, positions, width=120, height=30)
            _docx_energy_stripe_cell(cell, label, font_color, png, _CELL_W_CM, _CELL_H_CM)
        else:
            _docx_write_cell(cell, label, bold=True, size=5.8, color=font_color, fill=fill)

    # Leyenda candados
    lock_title = document.add_table(rows=1, cols=1)
    _docx_apply_table_grid(lock_title, [17.8])
    _docx_write_cell(lock_title.cell(0, 0), "Clasificación de Candados según sector y función", bold=True, size=7.0, color=MODO1_FONT_HEX, fill=MODO1_COLOR_HEX)
    locks = document.add_table(rows=1, cols=5)
    _docx_apply_table_grid(locks, [17.8 / 5] * 5)
    lock_items = [
        ("Mantenimiento\nIndustrial", "EF0000", "FFFFFF"),
        ("Calidad", "FFF200", "0F172A"),
        ("Producción", "16A34A", "FFFFFF"),
        ("Mantenimiento Edilicio\nContratistas", "0F7DBD", "FFFFFF"),
        ("Supervisor MMTO Industrial\n(Bloqueo Departamental)", "050505", "FFFFFF"),
    ]
    for idx, (label, fill, font_color) in enumerate(lock_items):
        _docx_write_cell(locks.cell(0, idx), label, bold=True, size=5.8, color=font_color, fill=fill)

    # LSR
    lsr_table = document.add_table(rows=1, cols=1)
    _docx_apply_table_grid(lsr_table, [17.8])
    lsr_cell = lsr_table.cell(0, 0)
    _docx_write_cell(lsr_cell, "LSR", bold=True, size=10.0, color="94A3B8", fill="FFFFFF")
    lsr_bytes, _ = _data_uri_to_bytes(lsr_uri)
    if lsr_bytes:
        try:
            _docx_clear_cell(lsr_cell)
            p = lsr_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(lsr_bytes), width=Cm(15.0))
            _docx_set_cell_borders(lsr_cell)
        except Exception:
            _docx_write_cell(lsr_cell, "LSR", bold=True, size=10.0, color="94A3B8", fill="FFFFFF")

    # Firmas
    footer = document.add_table(rows=1, cols=2)
    _docx_apply_table_grid(footer, [8.9, 8.9])
    _docx_write_cell(footer.cell(0, 0), f"Elaborado por: {elaborado_por or '-'}\nPuesto: {puesto_elaborado or '-'} · Fecha: {fecha_firma_txt}", size=6.4, fill="F8FAFC")
    _docx_write_cell(footer.cell(0, 1), f"Aprobado por: {aprobado_por or '-'}\nPuesto: {puesto_aprobado or '-'} · Fecha: {fecha_firma_txt}", size=6.4, fill="F8FAFC")

    for paragraph in document.paragraphs:
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


# ============================================================
# EXPORTACIÓN EXCEL
# ============================================================

def _xlsx_style_range(ws, cell_range, fill=None, font_color="0F172A", bold=False, size=8.0, align="center"):
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    side = Side(style="thin", color="111827")
    border = Border(left=side, right=side, top=side, bottom=side)
    for row in ws[cell_range]:
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
            cell.font = Font(name="Bahnschrift", size=size, bold=bold, color=font_color)
            if fill:
                cell.fill = PatternFill("solid", fgColor=fill.replace("#", ""))

def _xlsx_merge_write(ws, cell_range, value, *, fill=None, font_color="0F172A", bold=False, size=8.0, align="center"):
    top_left = cell_range.split(":", 1)[0]
    ws.merge_cells(cell_range)
    ws[top_left] = value
    _xlsx_style_range(ws, cell_range, fill=fill, font_color=font_color, bold=bold, size=size, align=align)

def _xlsx_add_image(ws, data_uri, anchor, *, max_width_px, max_height_px):
    if not data_uri:
        return
    try:
        from openpyxl.drawing.image import Image as XLImage
        from PIL import Image as PILImage
        image_bytes, _ = _data_uri_to_bytes(data_uri)
        if not image_bytes:
            return
        pil_img = PILImage.open(io.BytesIO(image_bytes)).convert("RGBA")
        pil_img.thumbnail((max_width_px, max_height_px), PILImage.LANCZOS)
        stream = io.BytesIO()
        pil_img.save(stream, format="PNG")
        stream.seek(0)
        xl_img = XLImage(stream)
        xl_img.anchor = anchor
        ws.add_image(xl_img)
    except Exception:
        return

def build_modo_2_excel_bytes(
    ctx: Dict[str, Any],
    *,
    codigo: str,
    revision: str,
    fecha: "datetime.date | str",
    organizacion: str,
    logo_uri: str = "",
    lsr_uri: str = "",
    personal_afectado: str,
    personal_autorizado: str,
    elaborado_por: str,
    aprobado_por: str,
    puesto_elaborado: str = "",
    puesto_aprobado: str = "",
    fecha_firma: "datetime.date | str | None" = None,
    eval_riesgos_codigo: str = "",
    eval_riesgos_fecha: str = "",
) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    fecha_txt = _fecha_export_txt(fecha)
    fecha_firma_txt = _fecha_export_txt(fecha_firma) or fecha_txt
    equipo_meta = _equipment_meta(ctx)

    wb = Workbook()
    ws = wb.active
    ws.title = "Modo 2"
    ws.sheet_view.showGridLines = False
    ws.page_setup.orientation = "portrait"
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_margins.left = 0.18
    ws.page_margins.right = 0.18
    ws.page_margins.top = 0.18
    ws.page_margins.bottom = 0.18
    ws.page_margins.header = 0
    ws.page_margins.footer = 0

    for col in range(1, 25):
        ws.column_dimensions[get_column_letter(col)].width = 4.2
    for row in range(1, 58):
        ws.row_dimensions[row].height = 18

    white = PatternFill("solid", fgColor="FFFFFF")
    for row in range(1, 58):
        for col in range(1, 25):
            ws.cell(row=row, column=col).fill = white
            ws.cell(row=row, column=col).font = Font(name="Bahnschrift", size=8)

    # Encabezado
    _xlsx_merge_write(ws, "A1:C3", "" if logo_uri else organizacion, fill="FFFFFF", font_color="B91C1C", bold=True, size=8)
    _xlsx_add_image(ws, logo_uri, "A1", max_width_px=78, max_height_px=52)
    _xlsx_merge_write(ws, "D1:P1", "CONTROL DE ENERGÍAS PELIGROSAS", fill="050505", font_color="FFFFFF", bold=True, size=9)
    _xlsx_merge_write(ws, "D2:P3", "PROCEDIMIENTO ESPECÍFICO PARA CONTROL DE ENERGÍAS", fill="FFFFFF", bold=True, size=10)
    for rng, label, value in (("Q1:S1", "Código", codigo), ("Q2:S2", "Revisión", revision), ("Q3:S3", "Fecha", fecha_txt)):
        _xlsx_merge_write(ws, rng, label, fill="E5E7EB", bold=True, size=8)
        value_rng = rng.replace("Q", "T").replace("S", "X")
        _xlsx_merge_write(ws, value_rng, value, fill="F8FAFC", bold=True, size=8)

    # Sitio + personal
    def _xlsx_write_label_value(ws, cell_range, label, value):
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.cell.rich_text import CellRichText, TextBlock
        from openpyxl.styles.fonts import Font as RFont
        top_left = cell_range.split(":", 1)[0]
        ws.merge_cells(cell_range)
        side = Side(style="thin", color="111827")
        border = Border(left=side, right=side, top=side, bottom=side)
        try:
            rich = CellRichText(
                TextBlock(RFont(name="Bahnschrift", size=8, bold=False, color="FF334155"), label + "\n"),
                TextBlock(RFont(name="Bahnschrift", size=10, bold=True, color="FF111827"), value),
            )
            ws[top_left] = rich
        except Exception:
            ws[top_left] = f"{label}\n{value}"
        cell = ws[top_left]
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
        cell.fill = PatternFill("solid", fgColor="FFFFFF")

    _xlsx_write_label_value(ws, "A4:C6", "Negocio:", ctx.get("negocio") or "-")
    _xlsx_write_label_value(ws, "D4:F6", "Sitio:", ctx.get("sitio") or "-")
    _xlsx_write_label_value(ws, "G4:I6", "Área:", ctx.get("area") or "-")
    _xlsx_write_label_value(ws, "J4:M6", "Línea:", ctx.get("linea") or "-")
    _xlsx_merge_write(ws, "N4:X4", "Personal afectado - Puestos de trabajo", fill=MODO1_COLOR_HEX, font_color=MODO1_FONT_HEX, bold=True, size=7)
    _xlsx_merge_write(ws, "N5:X5", personal_afectado, fill="FFFFFF", size=7)
    _xlsx_merge_write(ws, "N6:X6", "Personal autorizado - Puestos de trabajo", fill=MODO1_COLOR_HEX, font_color=MODO1_FONT_HEX, bold=True, size=7)
    _xlsx_merge_write(ws, "N7:X7", personal_autorizado, fill="FFFFFF", size=7)

    # Equipo
    _xlsx_merge_write(ws, "A8:C9", "Equipo:", fill="F8FAFC", bold=True, size=8)
    _xlsx_merge_write(ws, "D8:X9", f"{_upper(ctx.get('equipo') or 'EQUIPO')}\n{equipo_meta}", fill="FFFFFF", bold=True, size=11)

    # Evaluación de riesgos
    eval_ref_txt = "Evaluación de riesgos de referencia:  "
    if eval_riesgos_codigo:
        eval_ref_txt += f"Código: {eval_riesgos_codigo}"
    if eval_riesgos_fecha:
        eval_ref_txt += f"  ·  Fecha: {eval_riesgos_fecha}"
    _xlsx_merge_write(ws, "A10:X10", eval_ref_txt, fill="FFFFFF", font_color="0F172A", bold=True, size=7, align="left")

    # Tareas
    _xlsx_merge_write(ws, "A11:C11", "Puntos de\nBloqueo", fill=MODO1_COLOR_HEX, font_color=MODO1_FONT_HEX, bold=True, size=7)
    _xlsx_merge_write(ws, "D11:F11", "Modo de\nIntervención", fill=MODO1_COLOR_HEX, font_color=MODO1_FONT_HEX, bold=True, size=7)
    _xlsx_merge_write(ws, "G11:X11", "Listado de tareas aplicable al presente procedimiento", fill=MODO1_COLOR_HEX, font_color=MODO1_FONT_HEX, bold=True, size=7)
    _xlsx_merge_write(ws, "A12:C15", "2", fill="FFFFFF", bold=True, size=32)
    _xlsx_merge_write(ws, "D12:F15", "MODO 2", fill=MODO1_COLOR_HEX, font_color=MODO1_FONT_HEX, bold=True, size=14)
    _xlsx_merge_write(ws, "G12:X15", "\n".join(f"- {task}" for task in (ctx.get("tareas") or [])), fill="FFFFFF", size=7, align="left")

    # Barra + foto
    _xlsx_merge_write(ws, "A16:X16", "Procedimiento - Control de Energías Peligrosas", fill=MODO1_COLOR_HEX, font_color=MODO1_FONT_HEX, bold=True, size=8)
    _xlsx_merge_write(ws, "A17:X28", "FOTO DEL PASO A PASO", fill="FFFFFF", font_color="94A3B8", bold=True, size=14)
    _xlsx_add_image(ws, ctx.get("photo_uri", ""), "H17", max_width_px=520, max_height_px=210)

    # Procedimiento: Acción y Verificación
    _xlsx_merge_write(ws, "A29:L29", "Acción", fill="3B3B3B", font_color="FFFFFF", bold=True, size=7)
    _xlsx_merge_write(ws, "M29:X29", "Verificación", fill="3B3B3B", font_color="FFFFFF", bold=True, size=7)
    accion_txt = "\n\n".join(_ACCION_MODO2)
    verif_txt = "\n\n".join(_VERIFICACION_MODO2)
    _xlsx_merge_write(ws, "A30:L40", accion_txt, fill="FFFFFF", size=7.0, align="left")
    _xlsx_merge_write(ws, "M30:X40", verif_txt, fill="FFFFFF", size=7.0, align="left")

    # Leyendas energías
    _xlsx_merge_write(ws, "A41:X41", "Clasificación de Energías Peligrosas", fill=MODO1_COLOR_HEX, font_color=MODO1_FONT_HEX, bold=True, size=8)

    def _xlsx_energy_png(ws, row, col_start, col_end, label, font_color, bg, stripe_color, stripe_positions):
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
        side = Side(style="thin", color="111827")
        border = Border(left=side, right=side, top=side, bottom=side)
        col_letter_start = get_column_letter(col_start)
        col_letter_end = get_column_letter(col_end)
        rng = f"{col_letter_start}{row}:{col_letter_end}{row}"
        ws.merge_cells(rng)
        top_cell = ws.cell(row=row, column=col_start)
        top_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        top_cell.font = Font(name="Bahnschrift", size=7, bold=True, color=font_color.replace("#", ""))
        for c in range(col_start, col_end + 1):
            ws.cell(row=row, column=c).border = border
        if stripe_color is None:
            top_cell.fill = PatternFill("solid", fgColor=bg.replace("#", ""))
            top_cell.value = label
        else:
            top_cell.fill = PatternFill("solid", fgColor=bg.replace("#", ""))
            top_cell.value = label
            try:
                from PIL import Image as _PIL, ImageDraw as _Draw, ImageFont as _IFont
                W, H = 160, 32
                img = _PIL.new("RGB", (W, H), tuple(int(bg[i:i+2], 16) for i in (0, 2, 4)))
                draw = _Draw.Draw(img)
                sc_rgb = tuple(int(stripe_color[i:i+2], 16) for i in (0, 2, 4))
                for x0p, x1p in stripe_positions:
                    draw.rectangle([int(W * x0p / 100), 0, int(W * x1p / 100) - 1, H - 1], fill=sc_rgb)
                fc_rgb = tuple(int(font_color[i:i+2], 16) for i in (0, 2, 4))
                try:
                    font = _IFont.truetype("arial.ttf", 9)
                except Exception:
                    font = _IFont.load_default()
                bbox = draw.textbbox((0, 0), label, font=font)
                tw = bbox[2] - bbox[0]; th = bbox[3] - bbox[1]
                draw.text(((W - tw) / 2, (H - th) / 2 - bbox[1]), label, fill=fc_rgb, font=font)
                buf = io.BytesIO(); img.save(buf, format="PNG"); buf.seek(0)
                from openpyxl.drawing.image import Image as XLImage
                xl_img = XLImage(buf); xl_img.anchor = f"{col_letter_start}{row}"
                ws.add_image(xl_img)
            except Exception:
                pass

    _E_ROW1, _E_ROW2 = 42, 43
    xlsx_energy = [
        (_E_ROW1,  1,  4, "E: Eléctrica",     "FFFFFF", "000000", None, None),
        (_E_ROW1,  5,  8, "N: Neumática",      "FFFFFF", "0284C7", None, None),
        (_E_ROW1,  9, 12, "AM: Amoníaco",      "0F172A", "D9D9D9", "F59E0B", [(15, 26), (42, 53), (68, 79)]),
        (_E_ROW1, 13, 16, "T: Térmica",        "FFFFFF", "DC2626", None, None),
        (_E_ROW1, 17, 20, "H: Hidráulica",     "FFFFFF", "7C3AED", None, None),
        (_E_ROW1, 21, 24, "P: Potencial",      "0F172A", "FFF200", "050505", [(22, 34), (62, 74)]),
        (_E_ROW2,  1,  4, "Q: Química",        "0F172A", "FFF200", None, None),
        (_E_ROW2,  5,  8, "V: Vapor",          "111827", "F59E0B", None, None),
        (_E_ROW2,  9, 12, "A: Agua",           "FFFFFF", "16A34A", None, None),
        (_E_ROW2, 13, 16, "SC: Soda Cáustica", "0F172A", "D9D9D9", "F59E0B", [(25, 37), (60, 72)]),
        (_E_ROW2, 17, 20, "Oz: Ozono",         "0F172A", "BAE6FD", None, None),
        (_E_ROW2, 21, 24, "GC: Gas Carbónico", "111827", "C7D2FE", None, None),
    ]
    for entry in xlsx_energy:
        _xlsx_energy_png(ws, *entry)

    _xlsx_merge_write(ws, "A44:X44", "Clasificación de Candados según sector y función", fill=MODO1_COLOR_HEX, font_color=MODO1_FONT_HEX, bold=True, size=8)
    lock_items_xl = [
        ("A45:E46", "Mantenimiento\nIndustrial", "EF0000", "FFFFFF"),
        ("F45:J46", "Calidad", "FFF200", "0F172A"),
        ("K45:N46", "Producción", "16A34A", "FFFFFF"),
        ("O45:S46", "Mantenimiento Edilicio\nContratistas", "0F7DBD", "FFFFFF"),
        ("T45:X46", "Supervisor MMTO Industrial\n(Bloqueo Departamental)", "050505", "FFFFFF"),
    ]
    for rng, label, fill, font_color in lock_items_xl:
        _xlsx_merge_write(ws, rng, label, fill=fill, font_color=font_color, bold=True, size=7)

    # LSR
    _xlsx_merge_write(ws, "A47:X53", "LSR", fill="FFFFFF", font_color="94A3B8", bold=True, size=14)
    _xlsx_add_image(ws, lsr_uri, "D47", max_width_px=600, max_height_px=130)

    # Firmas
    _xlsx_merge_write(ws, "A54:L56", f"Elaborado por: {elaborado_por or '-'}\nPuesto: {puesto_elaborado or '-'} · Fecha: {fecha_firma_txt}", fill="F8FAFC", size=7)
    _xlsx_merge_write(ws, "M54:X56", f"Aprobado por: {aprobado_por or '-'}\nPuesto: {puesto_aprobado or '-'} · Fecha: {fecha_firma_txt}", fill="F8FAFC", size=7)

    for row in range(42, 44):
        ws.row_dimensions[row].height = 26
    for row in range(45, 47):
        ws.row_dimensions[row].height = 26
    for row in range(47, 54):
        ws.row_dimensions[row].height = 20
    for row in range(30, 41):
        ws.row_dimensions[row].height = 22

    ws.print_area = "A1:X56"
    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


# ============================================================
# INTERFAZ STREAMLIT
# ============================================================

st.markdown(
    """
    <style>
        html, body, [class*="css"], .stApp, input, textarea, button {
            font-family: Bahnschrift, 'Bahnschrift SemiCondensed', 'Arial Narrow', Arial, Helvetica, sans-serif !important;
        }
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("GENERADOR DE PROCEDIMIENTOS MODO 2")
st.markdown(
    """
    <div class="soft-note">
        Importá el borrador JSON generado por la app de análisis. Esta versión genera el procedimiento específico
        para <strong>Modo 2</strong> (ingreso de cuerpo completo o intervención con protecciones no entrelazadas), con diseño documental listo para descarga en Word DOCX y Excel XLSX.
    </div>
    """,
    unsafe_allow_html=True,
)

with st.sidebar:
    st.subheader("Configuración documental")
    codigo = st.text_input("Código", value="LOTO-M2-001")
    revision = st.text_input("Revisión", value="00")
    fecha = st.date_input("Fecha", value=datetime.date.today(), format="DD/MM/YYYY")
    organizacion = st.text_input("Organización / logo textual", value="ARCA CONTINENTAL")
    logo_file = st.file_uploader(
        "Logo opcional para reemplazar arca.png",
        type=["png", "jpg", "jpeg", "webp"],
        help="Si no se carga un archivo, el documento intenta usar arca.png desde la carpeta de la aplicación.",
    )

    st.divider()
    st.subheader("Firmas")
    elaborado_por = st.text_input("Elaborado por", value="")
    puesto_elaborado = st.text_input("Puesto de quien elabora", value="")
    aprobado_por = st.text_input("Aprobado por", value="")
    puesto_aprobado = st.text_input("Puesto de quien aprueba", value="")
    fecha_firma = st.date_input("Fecha de firmas", value=fecha, format="DD/MM/YYYY")

uploaded_json = st.file_uploader(
    "Importar borrador JSON",
    type=["json"],
    help="Usá el archivo generado por el botón 'Descargar borrador JSON' de la app anterior.",
)

if uploaded_json is None:
    st.info("Cargá un borrador JSON para generar el procedimiento.")
    st.stop()

try:
    payload = json.loads(uploaded_json.getvalue().decode("utf-8"))
except Exception as exc:
    st.error(f"No se pudo leer el JSON importado: {exc}")
    st.stop()

ctx_detected = extract_procedure_context(payload)

st.subheader("Datos importados editables")
st.caption("Todos los campos detectados desde el JSON pueden reescribirse antes de generar el procedimiento.")

id_col_1, id_col_2, id_col_3 = st.columns(3)
with id_col_1:
    negocio_edit = st.text_input("Negocio", value=_normalize(ctx_detected.get("negocio")), key="edit_negocio")
    sitio_edit = st.text_input("Sitio", value=_normalize(ctx_detected.get("sitio")), key="edit_sitio")
    area_edit = st.text_input("Área", value=_normalize(ctx_detected.get("area")), key="edit_area")
with id_col_2:
    linea_edit = st.text_input("Línea", value=_normalize(ctx_detected.get("linea")), key="edit_linea")
    equipo_edit = st.text_input("Equipo", value=_normalize(ctx_detected.get("equipo")), key="edit_equipo")
    fabricante_edit = st.text_input("Fabricante", value=_normalize(ctx_detected.get("fabricante")), key="edit_fabricante")
with id_col_3:
    modelo_edit = st.text_input("Modelo", value=_normalize(ctx_detected.get("modelo")), key="edit_modelo")
    anio_edit = st.text_input("Año", value=_normalize(ctx_detected.get("anio")), key="edit_anio")
    modo_inicial_edit = st.text_input("Modo inicial", value=_normalize(ctx_detected.get("modo_inicial")), key="edit_modo_inicial")

modo_final_edit = st.text_input("Modo final", value=_normalize(ctx_detected.get("modo_final")), key="edit_modo_final")

# Evaluación de riesgos
st.subheader("Evaluación de riesgos de referencia")
st.caption("Se completa automáticamente desde el JSON. Podés editarlo manualmente antes de generar.")
eval_col_1, eval_col_2 = st.columns(2)
with eval_col_1:
    eval_riesgos_codigo = st.text_input(
        "Código de la evaluación de riesgos",
        value=_normalize(ctx_detected.get("codigo_documento")),
        key="edit_eval_codigo",
        help="Se toma automáticamente del campo 'codigo_documento' del JSON importado.",
    )
with eval_col_2:
    eval_riesgos_fecha = st.text_input(
        "Fecha de la evaluación de riesgos",
        value=_normalize(ctx_detected.get("fecha_documento")),
        key="edit_eval_fecha",
        help="Se toma automáticamente del campo 'fecha_documento' del JSON importado.",
    )

st.subheader("Personal editable")
st.caption("Estos campos se escriben manualmente antes de generar el procedimiento y se imprimen en el encabezado del documento.")
person_col_1, person_col_2 = st.columns(2)
with person_col_1:
    personal_afectado = st.text_area(
        "Personal afectado - Puestos de trabajo",
        value="Técnicos del Sector de Mantenimiento.\nContratistas.",
        height=96,
        key="edit_personal_afectado",
    )
with person_col_2:
    personal_autorizado = st.text_area(
        "Personal autorizado - Puestos de trabajo",
        value="Técnicos del Sector Mantenimiento.\nContratistas.",
        height=96,
        key="edit_personal_autorizado",
    )

st.subheader("Tareas aplicables")
tasks_detected = ctx_detected.get("tareas") or []
tasks_text = st.text_area(
    "Listado de tareas (tareas_predefinidas + tareas_manuales combinadas)",
    value="\n".join(tasks_detected),
    height=140,
    help="Una tarea por línea. Se combinan automáticamente 'tareas_predefinidas' y 'tareas_manuales' del JSON, sin duplicados. Podés editar el resultado antes de generar.",
)

st.subheader("Foto del paso a paso")
step_photo_file = st.file_uploader(
    "Subir foto o imagen del paso a paso",
    type=["png", "jpg", "jpeg", "webp"],
    help="Esta imagen se inserta en el área central del procedimiento.",
)
photo_uri = _uploaded_file_data_uri(step_photo_file)
if step_photo_file is not None:
    st.image(step_photo_file, caption="Foto del paso a paso cargada", use_container_width=True)
else:
    st.caption("Si no se sube una imagen, el procedimiento mostrará un recuadro reservado para la foto del paso a paso.")

with st.expander("Lienzo opcional para tags sobre imagen"):
    st.caption("El lienzo se habilita solo si el entorno tiene instalada la extensión streamlit-drawable-canvas.")
    try:
        from streamlit_drawable_canvas import st_canvas
        from PIL import Image
        if step_photo_file is None:
            st.info("Subí primero una foto del paso a paso para activar el lienzo de marcado.")
        else:
            canvas_result = st_canvas(
                fill_color="rgba(255, 0, 0, 0.15)",
                stroke_width=3,
                stroke_color="#C00000",
                background_image=Image.open(step_photo_file),
                update_streamlit=True,
                height=360,
                drawing_mode="freedraw",
                key="step_photo_canvas",
            )
            if canvas_result.image_data is not None:
                annotated = Image.fromarray(canvas_result.image_data.astype("uint8"), "RGBA")
                buffer = __import__("io").BytesIO()
                annotated.save(buffer, format="PNG")
                photo_uri = _bytes_to_data_uri(buffer.getvalue(), "image/png")
                st.success("Se usará la imagen anotada del lienzo en el procedimiento.")
    except Exception:
        st.info("Lienzo no disponible en este entorno. Podés subir la foto ya marcada o continuar sin anotaciones.")

ctx = dict(ctx_detected)
ctx.update(
    {
        "negocio": negocio_edit,
        "sitio": sitio_edit,
        "area": area_edit,
        "linea": linea_edit,
        "equipo": equipo_edit,
        "fabricante": fabricante_edit,
        "modelo": modelo_edit,
        "anio": anio_edit,
        "modo_inicial": modo_inicial_edit,
        "modo_final": modo_final_edit,
        "tareas": _split_tasks(tasks_text),
        "photo_uri": photo_uri,
    }
)
modo_final = ctx.get("modo_final")

col_a, col_b, col_c = st.columns(3)
with col_a:
    st.metric("Modo final", _normalize(modo_final) or "Sin dato")
with col_b:
    st.metric("Equipo", _normalize(ctx.get("equipo"))[:32] or "Sin dato")
with col_c:
    st.metric("Tareas", len(ctx.get("tareas") or []))

if not _mode_is_modo_2(modo_final):
    st.markdown(
        f"""
        <div class="status-warn">
            El borrador importado no corresponde a <strong>Modo 2</strong>. Modo final detectado: <strong>{_html(modo_final or 'sin dato')}</strong>.
            Esta versión del generador solo habilita la plantilla Modo 2.
        </div>
        """,
        unsafe_allow_html=True,
    )
    with st.expander("Ver datos detectados del JSON"):
        st.json({"identificacion": ctx, "computed": payload.get("computed", {})})
    st.stop()

logo_uri = _logo_data_uri(logo_file)
lsr_uri = _lsr_data_uri()

procedure_html = build_modo_2_html(
    ctx,
    codigo=codigo,
    revision=revision,
    fecha=fecha,
    organizacion=organizacion,
    logo_uri=logo_uri,
    lsr_uri=lsr_uri,
    personal_afectado=personal_afectado,
    personal_autorizado=personal_autorizado,
    elaborado_por=elaborado_por,
    aprobado_por=aprobado_por,
    puesto_elaborado=puesto_elaborado,
    puesto_aprobado=puesto_aprobado,
    fecha_firma=fecha_firma,
    eval_riesgos_codigo=eval_riesgos_codigo,
    eval_riesgos_fecha=eval_riesgos_fecha,
)

st.markdown(
    """
    <div class="status-ok">
        Borrador compatible con <strong>Modo 2</strong>. Se generó la vista previa del procedimiento y quedaron habilitadas las descargas.
    </div>
    """,
    unsafe_allow_html=True,
)

st.subheader("Vista previa del procedimiento")
components.html(procedure_html, height=950, scrolling=True)

file_stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
base_filename = f"Procedimiento_Modo_2_{file_stamp}"

col_word, col_excel, col_html, col_json = st.columns(4)
with col_word:
    st.download_button(
        "Descargar Word DOCX",
        data=html_to_word_bytes(
            ctx,
            codigo=codigo,
            revision=revision,
            fecha=fecha,
            organizacion=organizacion,
            logo_uri=logo_uri,
            lsr_uri=lsr_uri,
            personal_afectado=personal_afectado,
            personal_autorizado=personal_autorizado,
            elaborado_por=elaborado_por,
            aprobado_por=aprobado_por,
            puesto_elaborado=puesto_elaborado,
            puesto_aprobado=puesto_aprobado,
            fecha_firma=fecha_firma,
            eval_riesgos_codigo=eval_riesgos_codigo,
            eval_riesgos_fecha=eval_riesgos_fecha,
        ),
        file_name=f"{base_filename}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

with col_excel:
    try:
        excel_bytes = build_modo_2_excel_bytes(
            ctx,
            codigo=codigo,
            revision=revision,
            fecha=fecha,
            organizacion=organizacion,
            logo_uri=logo_uri,
            lsr_uri=lsr_uri,
            personal_afectado=personal_afectado,
            personal_autorizado=personal_autorizado,
            elaborado_por=elaborado_por,
            aprobado_por=aprobado_por,
            puesto_elaborado=puesto_elaborado,
            puesto_aprobado=puesto_aprobado,
            fecha_firma=fecha_firma,
            eval_riesgos_codigo=eval_riesgos_codigo,
            eval_riesgos_fecha=eval_riesgos_fecha,
        )
        st.download_button(
            "Descargar Excel XLSX",
            data=excel_bytes,
            file_name=f"{base_filename}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
    except Exception as exc:
        st.warning(f"Excel editable no disponible en este entorno. Detalle técnico: {exc}")

with col_html:
    st.download_button(
        "Descargar HTML",
        data=procedure_html.encode("utf-8"),
        file_name=f"{base_filename}.html",
        mime="text/html",
        use_container_width=True,
    )

with col_json:
    st.download_button(
        "Descargar JSON importado",
        data=json.dumps(payload, ensure_ascii=False, indent=4).encode("utf-8"),
        file_name=f"Borrador_importado_{file_stamp}.json",
        mime="application/json",
        use_container_width=True,
    )

with st.expander("Datos editados para generación"):
    st.json(
        {
            "negocio": ctx.get("negocio"),
            "sitio": ctx.get("sitio"),
            "area": ctx.get("area"),
            "linea": ctx.get("linea"),
            "equipo": ctx.get("equipo"),
            "modo_inicial": ctx.get("modo_inicial"),
            "modo_final": ctx.get("modo_final"),
            "tareas": ctx.get("tareas"),
            "eval_riesgos_codigo": eval_riesgos_codigo,
            "eval_riesgos_fecha": eval_riesgos_fecha,
            "foto_paso_a_paso_cargada": bool(ctx.get("photo_uri")),
        }
    )
